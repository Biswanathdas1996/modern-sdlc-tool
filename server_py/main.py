import os
import re
import json
import base64
import asyncio
from io import BytesIO
from datetime import datetime
from typing import Optional, List, Dict, Any

from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, Request, Response, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, StreamingResponse
from sse_starlette.sse import EventSourceResponse
import httpx

from storage import storage
from models import AnalyzeRequest, RequirementsRequest
from ai import (
    analyze_repository, generate_documentation, generate_bpmn_diagram,
    generate_brd, generate_test_cases, generate_test_data,
    generate_user_stories, generate_copilot_prompt, find_related_stories,
    transcribe_audio
)
from jira_service import (
    sync_stories_to_jira, get_jira_stories, find_related_jira_stories,
    sync_subtask_to_jira, get_jira_parent_story_context
)
from agents.jira_agent import jira_agent
from agents.conversation_manager import conversation_manager
from mongodb_client import (
    ingest_document, search_knowledge_base, delete_document_chunks,
    get_knowledge_stats, create_knowledge_document_in_mongo,
    get_knowledge_documents_from_mongo, update_knowledge_document_in_mongo,
    delete_knowledge_document_from_mongo
)

app = FastAPI(title="DocuGen AI API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def log(message: str, source: str = "express"):
    formatted_time = datetime.now().strftime("%I:%M:%S %p")
    print(f"{formatted_time} [{source}] {message}")


@app.middleware("http")
async def log_requests(request: Request, call_next):
    start = datetime.now()
    response = await call_next(request)
    duration = (datetime.now() - start).total_seconds() * 1000
    
    if request.url.path.startswith("/api"):
        log(f"{request.method} {request.url.path} {response.status_code} in {duration:.0f}ms")
    
    return response


@app.get("/api/projects")
async def get_projects():
    try:
        projects = storage.get_all_projects()
        return [p.model_dump() for p in projects]
    except Exception as e:
        print(f"Error fetching projects: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch projects")


@app.get("/api/projects/{id}")
async def get_project(id: str):
    try:
        project = storage.get_project(id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        return project.model_dump()
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error fetching project: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch project")


@app.delete("/api/projects/{id}")
async def delete_project(id: str):
    try:
        project = storage.get_project(id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        storage.delete_project(id)
        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting project: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete project")


async def run_analysis(project_id: str, repo_url: str):
    try:
        analysis = await analyze_repository(repo_url, project_id)
        storage.create_analysis(analysis)
        
        tech_stack = []
        if analysis.get("techStack"):
            tech_stack = analysis["techStack"].get("languages", []) + analysis["techStack"].get("frameworks", [])
        
        storage.update_project(project_id, {
            "techStack": tech_stack,
            "description": analysis.get("summary", ""),
        })
        
        try:
            updated_project = storage.get_project(project_id)
            if updated_project:
                documentation = await generate_documentation(analysis, updated_project.model_dump())
                saved_doc = storage.create_documentation(documentation)
                
                try:
                    print("Generating BPMN diagrams for features...")
                    bpmn_data = await generate_bpmn_diagram(saved_doc.model_dump(), analysis)
                    storage.create_bpmn_diagram(bpmn_data)
                    print("BPMN diagrams generated successfully")
                except Exception as bpmn_error:
                    print(f"BPMN diagram generation error: {bpmn_error}")
        except Exception as doc_error:
            print(f"Documentation generation error: {doc_error}")
        
        storage.update_project(project_id, {"status": "completed"})
    except Exception as error:
        print(f"Analysis error: {error}")
        storage.update_project(project_id, {"status": "error"})


@app.post("/api/projects/analyze")
async def analyze_project(request: AnalyzeRequest, background_tasks: BackgroundTasks):
    try:
        repo_url = request.repoUrl
        
        match = re.match(r"github\.com/([^/]+)/([^/]+)", repo_url.replace("https://", ""))
        if not match:
            raise HTTPException(status_code=400, detail="Invalid GitHub URL")
        
        repo_name = f"{match.group(1)}/{match.group(2)}".replace(".git", "")
        
        project = storage.create_project({
            "name": repo_name,
            "repoUrl": repo_url,
            "techStack": [],
            "status": "analyzing",
        })
        
        background_tasks.add_task(run_analysis, project.id, repo_url)
        
        return JSONResponse(status_code=201, content=project.model_dump())
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error analyzing repository: {e}")
        raise HTTPException(status_code=500, detail="Failed to analyze repository")


@app.get("/api/analysis/current")
async def get_current_analysis():
    try:
        projects = storage.get_all_projects()
        if not projects:
            raise HTTPException(status_code=404, detail="No projects found")
        analysis = storage.get_analysis(projects[0].id)
        if not analysis:
            raise HTTPException(status_code=404, detail="No analysis found")
        return analysis.model_dump()
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error fetching analysis: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch analysis")


@app.get("/api/documentation/current")
async def get_current_documentation():
    try:
        projects = storage.get_all_projects()
        if not projects:
            raise HTTPException(status_code=404, detail="No projects found")
        doc = storage.get_documentation(projects[0].id)
        if not doc:
            raise HTTPException(status_code=404, detail="No documentation found")
        return doc.model_dump()
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error fetching documentation: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch documentation")


@app.get("/api/bpmn/current")
async def get_current_bpmn():
    try:
        projects = storage.get_all_projects()
        if not projects:
            raise HTTPException(status_code=404, detail="No projects found")
        doc = storage.get_documentation(projects[0].id)
        if not doc:
            raise HTTPException(status_code=404, detail="No documentation found")
        bpmn = storage.get_bpmn_diagram(doc.id)
        if not bpmn:
            raise HTTPException(status_code=404, detail="No BPMN diagrams found")
        return bpmn.model_dump()
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error fetching BPMN diagrams: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch BPMN diagrams")


@app.post("/api/bpmn/regenerate")
async def regenerate_bpmn():
    try:
        projects = storage.get_all_projects()
        if not projects:
            raise HTTPException(status_code=404, detail="No projects found")
        project = projects[0]
        doc = storage.get_documentation(project.id)
        if not doc:
            raise HTTPException(status_code=404, detail="No documentation found")
        analysis = storage.get_analysis(project.id)
        if not analysis:
            raise HTTPException(status_code=404, detail="No analysis found")
        
        storage.delete_bpmn_diagram(doc.id)
        bpmn_data = await generate_bpmn_diagram(doc.model_dump(), analysis.model_dump())
        new_bpmn = storage.create_bpmn_diagram(bpmn_data)
        
        return new_bpmn.model_dump()
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error regenerating BPMN diagram: {e}")
        raise HTTPException(status_code=500, detail="Failed to regenerate BPMN diagram")


@app.post("/api/database-schema/connect")
async def connect_database_schema(request: Request):
    try:
        import psycopg2
        
        body = await request.json()
        connection_string = body.get("connectionString", "").strip()
        
        if not connection_string:
            raise HTTPException(status_code=400, detail="Connection string is required")
        
        if connection_string.lower().startswith("psql "):
            connection_string = connection_string[5:].strip()
        if (connection_string.startswith("'") and connection_string.endswith("'")) or \
           (connection_string.startswith('"') and connection_string.endswith('"')):
            connection_string = connection_string[1:-1]
        
        projects = storage.get_all_projects()
        if not projects:
            raise HTTPException(status_code=400, detail="Please analyze a repository first")
        project = projects[0]
        
        try:
            conn = psycopg2.connect(connection_string)
            cursor = conn.cursor()
            
            cursor.execute("SELECT current_database()")
            result = cursor.fetchone()
            database_name = result[0] if result else "unknown"
            
            query = """
                SELECT 
                    t.table_name,
                    c.column_name,
                    c.data_type,
                    c.is_nullable,
                    c.column_default,
                    CASE WHEN pk.column_name IS NOT NULL THEN true ELSE false END as is_primary_key,
                    CASE WHEN fk.column_name IS NOT NULL THEN true ELSE false END as is_foreign_key,
                    fk.foreign_table_name || '.' || fk.foreign_column_name as references_column
                FROM information_schema.tables t
                JOIN information_schema.columns c ON t.table_name = c.table_name AND t.table_schema = c.table_schema
                LEFT JOIN (
                    SELECT ku.table_name, ku.column_name
                    FROM information_schema.table_constraints tc
                    JOIN information_schema.key_column_usage ku ON tc.constraint_name = ku.constraint_name
                    WHERE tc.constraint_type = 'PRIMARY KEY' AND tc.table_schema = 'public'
                ) pk ON c.table_name = pk.table_name AND c.column_name = pk.column_name
                LEFT JOIN (
                    SELECT 
                        kcu.table_name,
                        kcu.column_name,
                        ccu.table_name AS foreign_table_name,
                        ccu.column_name AS foreign_column_name
                    FROM information_schema.table_constraints tc
                    JOIN information_schema.key_column_usage kcu ON tc.constraint_name = kcu.constraint_name
                    JOIN information_schema.constraint_column_usage ccu ON tc.constraint_name = ccu.constraint_name
                    WHERE tc.constraint_type = 'FOREIGN KEY' AND tc.table_schema = 'public'
                ) fk ON c.table_name = fk.table_name AND c.column_name = fk.column_name
                WHERE t.table_schema = 'public' AND t.table_type = 'BASE TABLE'
                ORDER BY t.table_name, c.ordinal_position
            """
            
            cursor.execute(query)
            rows = cursor.fetchall()
            
            tables_map = {}
            for row in rows:
                table_name = row[0]
                if table_name not in tables_map:
                    tables_map[table_name] = []
                tables_map[table_name].append({
                    "name": row[1],
                    "dataType": row[2],
                    "isNullable": row[3] == "YES",
                    "defaultValue": row[4],
                    "isPrimaryKey": row[5],
                    "isForeignKey": row[6],
                    "references": row[7],
                })
            
            tables = []
            for table_name, columns in tables_map.items():
                cursor.execute(f'SELECT COUNT(*) FROM "{table_name}"')
                count_result = cursor.fetchone()
                row_count = count_result[0] if count_result else 0
                tables.append({
                    "name": table_name,
                    "columns": columns,
                    "rowCount": row_count,
                })
            
            cursor.close()
            conn.close()
            
            storage.delete_database_schema(project.id)
            
            masked_connection_string = re.sub(r"(://[^:]+:)[^@]+(@)", r"\1****\2", connection_string)
            
            schema_info = storage.create_database_schema({
                "projectId": project.id,
                "connectionString": masked_connection_string,
                "databaseName": database_name,
                "tables": tables,
            })
            
            documentation = storage.get_documentation(project.id)
            if documentation:
                storage.update_documentation(project.id, {
                    "databaseSchema": {
                        "databaseName": database_name,
                        "connectionString": masked_connection_string,
                        "tables": tables,
                    },
                })
                print(f"Database schema saved to documentation for project {project.id}")
            
            return schema_info.model_dump()
        except Exception as db_error:
            print(f"Database connection error: {db_error}")
            raise HTTPException(status_code=400, detail=f"Failed to connect to database: {str(db_error)}")
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error processing database schema: {e}")
        raise HTTPException(status_code=500, detail="Failed to process database schema")


@app.get("/api/database-schema/current")
async def get_current_database_schema():
    try:
        projects = storage.get_all_projects()
        if not projects:
            return None
        schema = storage.get_database_schema(projects[0].id)
        return schema.model_dump() if schema else None
    except Exception as e:
        print(f"Error fetching database schema: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch database schema")


@app.delete("/api/database-schema/current")
async def delete_current_database_schema():
    try:
        projects = storage.get_all_projects()
        if not projects:
            raise HTTPException(status_code=404, detail="No project found")
        project = projects[0]
        storage.delete_database_schema(project.id)
        
        documentation = storage.get_documentation(project.id)
        if documentation:
            storage.update_documentation(project.id, {"databaseSchema": None})
            print(f"Database schema removed from documentation for project {project.id}")
        
        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting database schema: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete database schema")


@app.post("/api/requirements")
async def create_requirements(
    title: str = Form(...),
    description: Optional[str] = Form(None),
    inputType: str = Form(...),
    requestType: str = Form("feature"),
    file: Optional[UploadFile] = File(None),
    audio: Optional[UploadFile] = File(None)
):
    try:
        projects = storage.get_all_projects()
        project_id = projects[0].id if projects else "global"
        
        final_description = description or ""
        
        if inputType == "audio" and audio:
            audio_buffer = await audio.read()
            final_description = await transcribe_audio(audio_buffer)
        
        if inputType == "file" and file:
            file_content = await file.read()
            final_description = file_content.decode("utf-8", errors="ignore")
        
        feature_request = storage.create_feature_request({
            "projectId": project_id,
            "title": title,
            "description": final_description,
            "inputType": inputType,
            "requestType": requestType,
            "rawInput": final_description,
        })
        
        return JSONResponse(status_code=201, content=feature_request.model_dump())
    except Exception as e:
        print(f"Error creating feature request: {e}")
        raise HTTPException(status_code=500, detail="Failed to create feature request")


@app.get("/api/brd/current")
async def get_current_brd():
    try:
        brd = storage.get_current_brd()
        if not brd:
            raise HTTPException(status_code=404, detail="No BRD found")
        return brd.model_dump()
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error fetching BRD: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch BRD")


@app.post("/api/brd/generate")
async def generate_brd_endpoint(request: Request):
    try:
        feature_request = storage.get_current_feature_request()
        if not feature_request:
            raise HTTPException(status_code=400, detail="No feature request found")
        
        projects = storage.get_all_projects()
        project_id = projects[0].id if projects else "global"
        
        analysis = storage.get_analysis(project_id) if projects else None
        documentation = storage.get_documentation(project_id) if projects else None
        database_schema = storage.get_database_schema(project_id) if projects else None
        
        knowledge_context = None
        knowledge_sources = []
        try:
            search_query = f"{feature_request.title} {feature_request.description}"
            kb_results = search_knowledge_base("global", search_query, 5)
            if kb_results:
                knowledge_context = "\n\n---\n\n".join([
                    f"[Source: {r['filename']}]\n{r['content']}" for r in kb_results
                ])
                knowledge_sources = [
                    {
                        "filename": r["filename"],
                        "chunkPreview": r["content"][:200] + ("..." if len(r["content"]) > 200 else "")
                    }
                    for r in kb_results
                ]
        except Exception as kb_error:
            print(f"Knowledge base search error: {kb_error}")
        
        doc_size = len(documentation.content) if documentation and documentation.content else 0
        kb_size = len(knowledge_context) if knowledge_context else 0
        feature_desc_size = len(feature_request.description) if feature_request.description else 0
        print(f"📋 BRD Generation Context:")
        print(f"  - Feature Request: '{feature_request.title}' ({feature_desc_size} chars)")
        print(f"  - Documentation: {'YES' if documentation else 'NO'} ({doc_size:,} chars)")
        print(f"  - Analysis: {'YES' if analysis else 'NO'}")
        print(f"  - DB Schema: {'YES' if database_schema else 'NO'}")
        print(f"  - Knowledge Base: {'YES' if knowledge_context else 'NO'} ({len(knowledge_sources)} chunks, {kb_size:,} chars)")
        
        async def generate():
            try:
                if knowledge_sources:
                    yield {"data": json.dumps({"knowledgeSources": knowledge_sources})}
                
                brd = await generate_brd(
                    feature_request.model_dump(),
                    analysis.model_dump() if analysis else None,
                    documentation.model_dump() if documentation else None,
                    database_schema.model_dump() if database_schema else None,
                    knowledge_context,
                    lambda chunk: None
                )
                
                brd["knowledgeSources"] = knowledge_sources if knowledge_sources else None
                storage.create_brd(brd)
                
                yield {"data": json.dumps({"content": json.dumps(brd.get("content", {}))})}
                yield {"data": json.dumps({"done": True})}
            except Exception as gen_error:
                print(f"BRD generation error: {gen_error}")
                yield {"data": json.dumps({"error": "Generation failed"})}
        
        return EventSourceResponse(generate())
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error generating BRD: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate BRD")


@app.get("/api/test-cases")
async def get_test_cases():
    try:
        brd = storage.get_current_brd()
        if not brd:
            return []
        test_cases = storage.get_test_cases(brd.id)
        return [tc.model_dump() for tc in test_cases]
    except Exception as e:
        print(f"Error fetching test cases: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch test cases")


@app.post("/api/test-cases/generate")
async def generate_test_cases_endpoint():
    try:
        brd = storage.get_current_brd()
        if not brd:
            raise HTTPException(status_code=400, detail="No BRD found. Please generate a BRD first.")
        
        projects = storage.get_all_projects()
        analysis = storage.get_analysis(projects[0].id) if projects else None
        documentation = storage.get_documentation(projects[0].id) if projects else None
        
        test_cases = await generate_test_cases(
            brd.model_dump(),
            analysis.model_dump() if analysis else None,
            documentation.model_dump() if documentation else None
        )
        
        if not test_cases:
            raise HTTPException(status_code=500, detail="Failed to generate test cases - no cases returned")
        
        saved_test_cases = storage.create_test_cases(test_cases)
        return [tc.model_dump() for tc in saved_test_cases]
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error generating test cases: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate test cases")


@app.get("/api/test-data")
async def get_test_data():
    try:
        test_data = storage.get_all_test_data()
        return [td.model_dump() for td in test_data]
    except Exception as e:
        print(f"Error fetching test data: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch test data")


@app.post("/api/test-data/generate")
async def generate_test_data_endpoint():
    try:
        brd = storage.get_current_brd()
        if not brd:
            raise HTTPException(status_code=400, detail="No BRD found. Please generate a BRD first.")
        
        test_cases = storage.get_test_cases(brd.id)
        if not test_cases:
            raise HTTPException(status_code=400, detail="No test cases found. Please generate test cases first.")
        
        projects = storage.get_all_projects()
        documentation = storage.get_documentation(projects[0].id) if projects else None
        
        test_data = await generate_test_data(
            [tc.model_dump() for tc in test_cases],
            brd.model_dump(),
            documentation.model_dump() if documentation else None
        )
        
        if not test_data:
            raise HTTPException(status_code=500, detail="Failed to generate test data - no data returned")
        
        saved_test_data = storage.create_test_data_batch(test_data)
        return [td.model_dump() for td in saved_test_data]
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error generating test data: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate test data")


@app.get("/api/user-stories/{brd_id}")
async def get_user_stories(brd_id: str):
    try:
        user_stories = storage.get_user_stories(brd_id)
        return [s.model_dump() for s in user_stories]
    except Exception as e:
        print(f"Error fetching user stories: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch user stories")


@app.post("/api/user-stories/generate")
async def generate_user_stories_endpoint(request: Request):
    try:
        body = await request.json() if request.headers.get("content-type") == "application/json" else {}
        parent_jira_key = body.get("parentJiraKey")
        
        brd = storage.get_current_brd()
        if not brd and body.get("brdData"):
            from models import BRD as BRDModel, BRDContent as BRDContentModel
            brd_data = body["brdData"]
            try:
                content = brd_data.get("content", {})
                if isinstance(content, dict):
                    brd_content = BRDContentModel(**content)
                else:
                    brd_content = content
                now = datetime.now().isoformat()
                brd = BRDModel(
                    id=brd_data.get("id", "restored"),
                    projectId=brd_data.get("projectId", ""),
                    featureRequestId=brd_data.get("featureRequestId", ""),
                    requestType=brd_data.get("requestType", "feature"),
                    title=brd_data.get("title", ""),
                    version=brd_data.get("version", "1.0"),
                    status=brd_data.get("status", "draft"),
                    content=brd_content,
                    knowledgeSources=brd_data.get("knowledgeSources", []),
                    createdAt=brd_data.get("createdAt", now),
                    updatedAt=brd_data.get("updatedAt", now),
                )
                storage.brds[brd.id] = brd
                storage.current_brd_id = brd.id
                print(f"Restored BRD from request body: {brd.title}")
            except Exception as restore_err:
                print(f"Error restoring BRD from body: {restore_err}")
        
        if not brd:
            raise HTTPException(status_code=400, detail="No BRD found. Please generate a BRD first.")
        
        projects = storage.get_all_projects()
        project_id = projects[0].id if projects else "global"
        documentation = storage.get_documentation(project_id) if projects else None
        database_schema = storage.get_database_schema(project_id) if projects else None
        
        parent_context = None
        if parent_jira_key:
            parent_context = await get_jira_parent_story_context(parent_jira_key)
        
        knowledge_context = None
        try:
            search_query = f"{brd.title} {brd.content.overview if hasattr(brd.content, 'overview') else ''}"
            kb_results = search_knowledge_base("global", search_query, 5)
            if kb_results:
                knowledge_context = "\n\n---\n\n".join([
                    f"[Source: {r['filename']}]\n{r['content']}" for r in kb_results
                ])
        except Exception as kb_error:
            print(f"Knowledge base search error: {kb_error}")
        
        user_stories = await generate_user_stories(
            brd.model_dump(),
            documentation.model_dump() if documentation else None,
            database_schema.model_dump() if database_schema else None,
            parent_context,
            knowledge_context
        )
        
        if not user_stories:
            raise HTTPException(status_code=500, detail="Failed to generate user stories - no stories returned")
        
        stories_with_parent = [
            {**story, "parentJiraKey": parent_jira_key}
            for story in user_stories
        ]
        
        saved_stories = storage.create_user_stories(stories_with_parent)
        return [s.model_dump() for s in saved_stories]
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error generating user stories: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate user stories")


@app.patch("/api/user-stories/{id}")
async def update_user_story(id: str, request: Request):
    try:
        updates = await request.json()
        updated_story = storage.update_user_story(id, updates)
        if not updated_story:
            raise HTTPException(status_code=404, detail="User story not found")
        return updated_story.model_dump()
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error updating user story: {e}")
        raise HTTPException(status_code=500, detail="Failed to update user story")


@app.delete("/api/user-stories/{id}")
async def delete_user_story(id: str):
    try:
        deleted = storage.delete_user_story(id)
        if not deleted:
            raise HTTPException(status_code=404, detail="User story not found")
        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting user story: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete user story")


@app.post("/api/copilot-prompt/generate")
async def generate_copilot_prompt_endpoint():
    try:
        brd = storage.get_current_brd()
        if not brd:
            raise HTTPException(status_code=400, detail="No BRD found. Please generate a BRD first.")
        
        user_stories = storage.get_user_stories(brd.id)
        if not user_stories:
            raise HTTPException(status_code=400, detail="No user stories found. Please generate user stories first.")
        
        projects = storage.get_all_projects()
        documentation = storage.get_documentation(projects[0].id) if projects else None
        analysis = storage.get_analysis(projects[0].id) if projects else None
        database_schema = storage.get_database_schema(projects[0].id) if projects else None
        
        prompt = await generate_copilot_prompt(
            [s.model_dump() for s in user_stories],
            documentation.model_dump() if documentation else None,
            analysis.model_dump() if analysis else None,
            database_schema.model_dump() if database_schema else None
        )
        
        return {"prompt": prompt}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error generating Copilot prompt: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate Copilot prompt")


@app.post("/api/jira/sync")
async def sync_to_jira():
    try:
        brd = storage.get_current_brd()
        if not brd:
            raise HTTPException(status_code=400, detail="No BRD found. Please generate a BRD first.")
        
        user_stories = storage.get_user_stories(brd.id)
        if not user_stories:
            raise HTTPException(status_code=400, detail="No user stories found. Please generate user stories first.")
        
        result = await sync_stories_to_jira(user_stories, storage)
        return result
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error syncing to JIRA: {e}")
        raise HTTPException(status_code=500, detail="Failed to sync to JIRA")


@app.get("/api/jira/stories")
async def get_jira_stories_endpoint():
    try:
        stories = await get_jira_stories()
        return stories
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error fetching JIRA stories: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch JIRA stories")


@app.post("/api/jira/find-related")
async def find_related_jira_stories_endpoint(request: Request):
    try:
        body = await request.json()
        feature_description = body.get("featureDescription")
        if not feature_description:
            raise HTTPException(status_code=400, detail="Feature description is required")
        
        related_stories = await find_related_jira_stories(feature_description)
        return {"relatedStories": related_stories}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error finding related stories: {e}")
        raise HTTPException(status_code=500, detail="Failed to find related stories")


@app.post("/api/v1/jira-agent/chat")
async def chat_with_jira_agent(request: Request):
    """Interactive chat endpoint with smart information gathering."""
    try:
        import uuid
        body = await request.json()
        prompt = body.get("prompt")
        session_id = body.get("session_id") or str(uuid.uuid4())
        context_data = body.get("context_data", {})
        
        if not prompt:
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "error": "Prompt is required",
                    "response": "Please provide a query prompt"
                }
            )
        
        # Get or create conversation context with memory
        conversation_ctx = conversation_manager.get_or_create_context(session_id)
        conversation_ctx.add_message("user", prompt)
        
        # Log conversation context
        print(f"📝 Session {session_id}: {len(conversation_ctx.messages)} messages in history")
        
        # Process with conversation context and memory
        result = await jira_agent.process_query_interactive(
            prompt,
            conversation_ctx=conversation_ctx,
            context_data=context_data
        )
        
        # Add agent response to conversation memory
        conversation_ctx.add_message("assistant", result.get("response", ""))
        
        return JSONResponse(
            content={
                "success": result.get("success", False),
                "session_id": session_id,
                "state": result.get("state", "initial"),
                "response": result.get("response", ""),
                "missing_fields": result.get("missing_fields", []),
                "tickets": result.get("tickets"),
                "intent": result.get("intent"),
                "error": result.get("error"),
                "collected_data": result.get("collected_data"),
                # Include conversation metadata
                "conversation_info": {
                    "message_count": len(conversation_ctx.messages),
                    "summary": conversation_ctx.get_summary(),
                    "state": conversation_ctx.state.value
                }
            }
        )
    except Exception as e:
        print(f"Error in JIRA agent chat: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "response": f"Failed to process chat request: {str(e)}"
            }
        )


@app.delete("/api/v1/jira-agent/session/{session_id}")
async def end_jira_agent_session(session_id: str):
    """End a conversation session and clear its context."""
    try:
        conversation_manager.delete_context(session_id)
        return {
            "message": "Session ended successfully",
            "session_id": session_id
        }
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to end session: {str(e)}"}
        )


@app.get("/api/v1/jira-agent/session/{session_id}")
async def get_jira_agent_session(session_id: str):
    """Get conversation history and context for a session."""
    try:
        context = conversation_manager.get_context(session_id)
        
        if not context:
            return JSONResponse(
                status_code=404,
                content={
                    "error": "Session not found",
                    "session_id": session_id
                }
            )
        
        return {
            "session_id": session_id,
            "state": context.state.value,
            "summary": context.get_summary(),
            "message_count": len(context.messages),
            "messages": context.messages,
            "collected_data": context.collected_data,
            "created_at": context.created_at.isoformat(),
            "updated_at": context.updated_at.isoformat()
        }
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to get session: {str(e)}"}
        )


@app.post("/api/v1/jira-agent/process")
async def process_jira_query_with_agent(request: Request):
    """Process natural language JIRA query using the AI agent."""
    try:
        body = await request.json()
        prompt = body.get("prompt")
        
        if not prompt:
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "error": "Prompt is required",
                    "response": "Please provide a query prompt"
                }
            )
        
        result = await jira_agent.process_query(prompt)
        
        return JSONResponse(
            content={
                "success": result.get("success", False),
                "prompt": result.get("prompt", ""),
                "intent": result.get("intent"),
                "response": result.get("response", ""),
                "tickets": result.get("tickets", []),
                "error": result.get("error")
            }
        )
    except Exception as e:
        print(f"Error in JIRA agent process: {e}")
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "response": f"Failed to process JIRA query: {str(e)}"
            }
        )


@app.get("/api/v1/jira-agent/health")
async def jira_agent_health():
    """Health check for JIRA agent."""
    active_conversations = conversation_manager.get_active_count()
    return {
        "status": "healthy", 
        "service": "jira-agent",
        "capabilities": ["intelligent_query_processing", "search", "create", "update", "chained_operations", "interactive_chat"],
        "features": {
            "interactive_mode": True,
            "smart_info_gathering": True,
            "multi_turn_conversations": True,
            "active_conversations": active_conversations
        }
    }


@app.post("/api/jira/sync-subtask")
async def sync_subtask_to_jira_endpoint(request: Request):
    try:
        body = await request.json()
        story_id = body.get("storyId")
        parent_key = body.get("parentKey")
        
        if not story_id or not parent_key:
            raise HTTPException(status_code=400, detail="Story ID and parent JIRA key are required")
        
        brd = storage.get_current_brd()
        if not brd:
            raise HTTPException(status_code=400, detail="No BRD found")
        
        user_stories = storage.get_user_stories(brd.id)
        story = next((s for s in user_stories if s.id == story_id), None)
        
        if not story:
            raise HTTPException(status_code=404, detail="User story not found")
        
        result = await sync_subtask_to_jira(story, parent_key, storage)
        return result
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error creating JIRA subtask: {e}")
        raise HTTPException(status_code=500, detail="Failed to create JIRA subtask")


@app.post("/api/confluence/publish")
async def publish_to_confluence(request: Request):
    try:
        body = await request.json()
        brd_id = body.get("brdId")
        
        jira_email = os.environ.get("JIRA_EMAIL")
        jira_token = os.environ.get("JIRA_API_TOKEN")
        confluence_instance_url = os.environ.get("JIRA_INSTANCE_URL", "daspapun21.atlassian.net")
        confluence_space_key = os.environ.get("CONFLUENCE_SPACE_KEY", "~5caf6d452c573b4b24d0f933")
        
        if not jira_email or not jira_token:
            raise HTTPException(status_code=400, detail="Confluence credentials not configured.")
        
        brd = storage.get_brd(brd_id) if brd_id else storage.get_current_brd()
        if not brd:
            raise HTTPException(status_code=400, detail="No BRD found. Please generate a BRD first.")
        
        auth = base64.b64encode(f"{jira_email}:{jira_token}".encode()).decode()
        confluence_base_url = f"https://{confluence_instance_url}/wiki/api/v2"
        
        adf_content = build_confluence_content(brd.model_dump())
        
        async with httpx.AsyncClient() as client:
            space_response = await client.get(
                f"https://{confluence_instance_url}/wiki/api/v2/spaces",
                params={"keys": confluence_space_key},
                headers={"Authorization": f"Basic {auth}", "Accept": "application/json"}
            )
            
            space_id = confluence_space_key
            if space_response.status_code == 200:
                space_data = space_response.json()
                if space_data.get("results"):
                    space_id = space_data["results"][0].get("id", confluence_space_key)
            
            create_page_body = {
                "spaceId": space_id,
                "status": "current",
                "title": f"BRD: {brd.title} - {datetime.now().strftime('%Y-%m-%d')}",
                "body": {
                    "representation": "atlas_doc_format",
                    "value": json.dumps(adf_content)
                }
            }
            
            response = await client.post(
                f"{confluence_base_url}/pages",
                headers={
                    "Authorization": f"Basic {auth}",
                    "Accept": "application/json",
                    "Content-Type": "application/json"
                },
                json=create_page_body
            )
            
            if response.status_code == 200 or response.status_code == 201:
                data = response.json()
                page_id = data.get("id", "")
                webui_link = data.get("_links", {}).get("webui", "")
                if webui_link:
                    page_url = f"https://{confluence_instance_url}/wiki{webui_link}"
                else:
                    page_url = f"https://{confluence_instance_url}/wiki/spaces/{confluence_space_key}/pages/{page_id}"
                return {
                    "success": True,
                    "pageId": page_id,
                    "pageUrl": page_url,
                    "message": "BRD published to Confluence successfully"
                }
            else:
                print(f"Confluence API error ({response.status_code}): {response.text}")
                raise HTTPException(status_code=response.status_code, detail=f"Failed to publish to Confluence: {response.status_code}")
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error publishing to Confluence: {e}")
        raise HTTPException(status_code=500, detail="Failed to publish to Confluence")


def build_confluence_content(brd: Dict[str, Any]) -> Dict[str, Any]:
    content = brd.get("content", {})
    
    def create_bullet_list(items: List[str]) -> Dict[str, Any]:
        if not items:
            return {"type": "paragraph", "content": [{"type": "text", "text": "None specified"}]}
        return {
            "type": "bulletList",
            "content": [
                {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": item or "N/A"}]}]}
                for item in items
            ]
        }
    
    return {
        "type": "doc",
        "version": 1,
        "content": [
            {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": brd.get("title", "")}]},
            {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "Overview"}]},
            {"type": "paragraph", "content": [{"type": "text", "text": content.get("overview", "No overview provided")}]},
            {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "Objectives"}]},
            create_bullet_list(content.get("objectives", [])),
            {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "Scope"}]},
            {"type": "heading", "attrs": {"level": 3}, "content": [{"type": "text", "text": "In Scope"}]},
            create_bullet_list(content.get("scope", {}).get("inScope", [])),
            {"type": "heading", "attrs": {"level": 3}, "content": [{"type": "text", "text": "Out of Scope"}]},
            create_bullet_list(content.get("scope", {}).get("outOfScope", [])),
            {"type": "rule"},
            {"type": "paragraph", "content": [{"type": "text", "text": f"Generated by DocuGen AI | Version: {brd.get('version', '1.0')} | Status: {brd.get('status', 'draft')}", "marks": [{"type": "em"}]}]}
        ]
    }


@app.get("/api/knowledge-base")
async def get_knowledge_base():
    try:
        documents = get_knowledge_documents_from_mongo()
        return documents
    except Exception as e:
        print(f"Error fetching knowledge documents: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch knowledge documents")


@app.get("/api/knowledge-base/stats")
async def get_knowledge_base_stats():
    try:
        stats = get_knowledge_stats("global")
        return stats
    except Exception as e:
        print(f"Error fetching knowledge stats: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch knowledge stats")


@app.post("/api/knowledge-base/upload")
async def upload_knowledge_document(file: UploadFile = File(...)):
    try:
        if not file:
            raise HTTPException(status_code=400, detail="No file provided")
        
        project_id = "global"
        
        file_content = await file.read()
        content = ""
        
        try:
            if file.content_type in ["text/plain", "text/markdown", "text/csv"]:
                content = file_content.decode("utf-8", errors="ignore")
            elif file.content_type == "application/json":
                json_content = json.loads(file_content.decode("utf-8"))
                content = json.dumps(json_content, indent=2)
            elif file.content_type == "application/pdf":
                from PyPDF2 import PdfReader
                reader = PdfReader(BytesIO(file_content))
                content = "\n".join([page.extract_text() or "" for page in reader.pages])
            elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                from docx import Document
                doc = Document(BytesIO(file_content))
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                content = file_content.decode("utf-8", errors="ignore")
        except Exception as parse_error:
            print(f"Error parsing file: {parse_error}")
            content = file_content.decode("utf-8", errors="ignore")
        
        if not content or len(content.strip()) < 50:
            raise HTTPException(status_code=400, detail="File content too short or could not be extracted")
        
        doc = create_knowledge_document_in_mongo({
            "projectId": project_id,
            "filename": file.filename or "unknown",
            "originalName": file.filename or "unknown",
            "contentType": file.content_type or "text/plain",
            "size": len(file_content),
        })
        
        try:
            chunk_count = ingest_document(doc["id"], project_id, doc["filename"], content)
            update_knowledge_document_in_mongo(doc["id"], {"chunkCount": chunk_count, "status": "ready"})
            doc["chunkCount"] = chunk_count
            doc["status"] = "ready"
        except Exception as ingest_error:
            print(f"Error ingesting document: {ingest_error}")
            update_knowledge_document_in_mongo(doc["id"], {"status": "error", "errorMessage": str(ingest_error)})
            doc["status"] = "error"
            doc["errorMessage"] = str(ingest_error)
        
        return JSONResponse(status_code=201, content=doc)
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error uploading knowledge document: {e}")
        raise HTTPException(status_code=500, detail="Failed to upload knowledge document")


@app.delete("/api/knowledge-base/{id}")
async def delete_knowledge_document(id: str):
    try:
        delete_document_chunks(id)
        deleted = delete_knowledge_document_from_mongo(id)
        if not deleted:
            raise HTTPException(status_code=404, detail="Document not found")
        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting knowledge document: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete knowledge document")


@app.post("/api/knowledge-base/search")
async def search_knowledge(request: Request):
    try:
        body = await request.json()
        query = body.get("query", "")
        limit = body.get("limit", 5)
        
        results = search_knowledge_base("global", query, limit)
        return results
    except Exception as e:
        print(f"Error searching knowledge base: {e}")
        raise HTTPException(status_code=500, detail="Failed to search knowledge base")


VITE_DEV_SERVER = "http://localhost:5173"


@app.api_route("/{path:path}", methods=["GET", "POST", "PUT", "DELETE", "PATCH", "OPTIONS"])
async def proxy_to_vite(request: Request, path: str):
    if path.startswith("api/"):
        raise HTTPException(status_code=404, detail="API endpoint not found")
    
    if os.environ.get("NODE_ENV") == "production":
        try:
            # Production assets are in dist/public/ (built by Vite)
            dist_path = "../dist/public"
            static_path = f"{dist_path}/{path}" if path else f"{dist_path}/index.html"
            
            # Handle assets subdirectory
            if path.startswith("assets/") or os.path.exists(static_path):
                if os.path.exists(static_path) and os.path.isfile(static_path):
                    with open(static_path, "rb") as f:
                        content = f.read()
                    content_type = "text/html"
                    if path.endswith(".js"):
                        content_type = "application/javascript"
                    elif path.endswith(".css"):
                        content_type = "text/css"
                    elif path.endswith(".json"):
                        content_type = "application/json"
                    elif path.endswith(".png"):
                        content_type = "image/png"
                    elif path.endswith(".svg"):
                        content_type = "image/svg+xml"
                    elif path.endswith(".ico"):
                        content_type = "image/x-icon"
                    elif path.endswith(".woff2"):
                        content_type = "font/woff2"
                    elif path.endswith(".woff"):
                        content_type = "font/woff"
                    return Response(content=content, media_type=content_type)
            
            # SPA fallback - serve index.html for all other routes
            with open(f"{dist_path}/index.html", "rb") as f:
                content = f.read()
            return Response(content=content, media_type="text/html")
        except Exception as e:
            print(f"Static file error: {e}")
            raise HTTPException(status_code=404, detail="Not found")
    else:
        try:
            async with httpx.AsyncClient() as client:
                url = f"{VITE_DEV_SERVER}/{path}"
                if request.query_params:
                    url += f"?{request.query_params}"
                
                response = await client.request(
                    method=request.method,
                    url=url,
                    headers={k: v for k, v in request.headers.items() if k.lower() not in ["host", "content-length"]},
                    content=await request.body() if request.method in ["POST", "PUT", "PATCH"] else None,
                    timeout=30.0
                )
                
                return Response(
                    content=response.content,
                    status_code=response.status_code,
                    headers={k: v for k, v in response.headers.items() if k.lower() not in ["transfer-encoding", "content-encoding"]}
                )
        except Exception as e:
            return Response(
                content=f"<!DOCTYPE html><html><body><h1>Vite Dev Server Not Running</h1><p>Please start Vite: npm run dev:client</p><p>Error: {e}</p></body></html>",
                media_type="text/html"
            )


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", "5000"))
    uvicorn.run(app, host="0.0.0.0", port=port)
